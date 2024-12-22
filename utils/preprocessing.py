import asyncio
import pymupdf
import docx
from io import BytesIO
from config_reader import config
from typing import List, Dict, Union


async def read_pdf(file: BytesIO) -> str:
    """
    Extracts text from pdf
    Arguments
    ---------
    * file - BytesIO object

    Returns
    -------
    * String with full document text    
    """
    # with pymupdf.open('pdf', file) as pdf:
    #     return chr(12).join([page.get_text() for page in pdf])
    try:
        # Open the PDF and extract text using asyncio.to_thread
        return await asyncio.to_thread(
            lambda: "\f".join([page.get_text() for page in pymupdf.open(stream=file, filetype="pdf")])
        )
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {e}")

async def read_docx(file: BytesIO) -> str:
    """
    Extracts text from docx
    Arguments
    ---------
    * file - BytesIO object

    Returns
    -------
    * String with full document text    
    """
    # doc = docx.Document(file)
    # fullText = []
    # for para in doc.paragraphs:
    #     fullText.append(para.text)
    # return '\n'.join(fullText)
    try:
        # Open the DOCX and extract paragraphs using asyncio.to_thread
        doc = await asyncio.to_thread(docx.Document, file)
        return await asyncio.to_thread(lambda: "\n".join(para.text for para in doc.paragraphs))
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {e}")